import os
import shutil
import streamlit as st
from unstructured.partition.docx import partition_docx
from unstructured.partition.pdf import partition_pdf
from elasticsearch import Elasticsearch
from langchain_core.documents import Document
from langchain_text_splitters import CharacterTextSplitter
import fitz  # PyMuPDF for PDF text extraction
import docx2txt  # For extracting text from DOCX files
import docx  # For creating DOCX files
from elasticsearch import Elasticsearch
from transformers import AutoModel  # Hugging Face model for embeddings
import re  # Regular expressions for text cleaning
import nltk  # Natural Language Toolkit for NLP tasks
import asyncio  # Asynchronous programming
from googletrans import Translator  # For text translation

# Download necessary NLTK data for natural language processing tasks
# 'punkt' is used for tokenizing text into sentences or words
# 'averaged_perceptron_tagger' is used for part-of-speech tagging
nltk.download('punkt')
nltk.download('averaged_perceptron_tagger')

# Initialize embeddings model from Hugging Face. This model will be used to generate embeddings for document chunks.
# Using trust_remote_code=True allows us to use the encoding function for our text inputs.
os.environ["KMP_DUPLICATE_LIB_OK"] = "TRUE"
os.environ["TOKENIZERS_PARALLELISM"] = "FALSE" 
embeddings_client = AutoModel.from_pretrained('jinaai/jina-embeddings-v2-base-en', trust_remote_code=True)

# Elasticsearch index name for storing document embeddings and metadata
INDEX_NAME = "chatbot_knn_translate"

# Initialize Elasticsearch client using API key for authentication. This connects to the Elasticsearch service.
es = Elasticsearch(
    "https://1c05a00ae549470bb2bff458f27505f2.us-central1.gcp.cloud.es.io:443",
    api_key="MGJTTGVKVUJ1STR4MElrOTVlYkI6anN6YnJNeG1RbHFkOUNVdEZQVGFtUQ=="
)

# Function to create the Elasticsearch index with kNN settings for similarity search
def create_index():
    if not es.indices.exists(index=INDEX_NAME):
        es.indices.create(index=INDEX_NAME, body={
            "mappings": {
                "properties": {
                    "embedding": {"type": "dense_vector", "dims": 768, "index": True, "similarity": "cosine"},
                    "page_content": {"type": "text"},
                    "metadata": {
                        "properties": {
                            "chunk_id": {"type": "keyword"},
                            "filename": {"type": "keyword"},
                            "page_number": {"type": "keyword"}
                        }
                    }
                }
            }
        })

# Ensure the index is created at the start of the application
create_index()

# Folder to store uploaded files temporarily
TEMP_FOLDER = "temp_knn"
UPLOADS_FOLDER = "uploads_translate_knn"
os.makedirs(TEMP_FOLDER, exist_ok=True)
os.makedirs(UPLOADS_FOLDER, exist_ok=True)

# Initialize the text splitter for dividing documents into manageable chunks
text_splitter = CharacterTextSplitter(
    separator="\n\n",
    chunk_size=3000,
    chunk_overlap=0,
    length_function=len,
    is_separator_regex=False,
)

# Function to save uploaded files and translate their content
async def save_uploaded_file(uploaded_file):
    # Save the uploaded file to the temporary folder
    if uploaded_file.name.endswith(".pdf"):
        file_path = os.path.join(TEMP_FOLDER, uploaded_file.name)
    with open(file_path, "wb") as f:
        f.write(uploaded_file.getvalue())

    # Extract text from the uploaded file
    if file_path.endswith(".pdf"):
        doc = fitz.open(file_path)
        text = "\n".join([page.get_text() for page in doc])
    elif file_path.endswith(".docx"):
        text = docx2txt.process(file_path)

    # Translate the extracted text to English
    translated_text = ""
    texts = text_splitter.split_text(text)
    async with Translator() as translator:
        translations = await translator.translate(texts, dest='en')
        for translation in translations:
            translated_text += translation.text

    # Clean the translated text by removing non-printable characters
    clean_text = re.sub(r'[\x00-\x1F\x7F-\x9F]', '', translated_text)

    # Save the translated text as a new DOCX file
    if uploaded_file.name.endswith(".pdf"):
        file_name = uploaded_file.name.replace(".pdf", ".docx")
    new_file_path = os.path.join(UPLOADS_FOLDER, file_name)
    document = docx.Document()
    document.add_paragraph(clean_text)
    document.save(new_file_path)

    return file_path, new_file_path

# Function to check if a document is already indexed in Elasticsearch
def is_document_indexed(file_name):
    query = {"query": {"match": {"metadata.filename": file_name}}}
    response = es.search(index=INDEX_NAME, body=query)
    return response['hits']['total']['value'] > 0

# Function to chunk documents into smaller sections for indexing
def document_retrieval_chunking(new_file_path):
    try:
        if new_file_path.endswith(".pdf"):
            return partition_pdf(
                filename=new_file_path,
                infer_table_structure=True,
                strategy="hi_res",
                extract_image_block_types=["Image"],
                chunking_strategy="by_title",
                max_characters=10000,
                combine_text_under_n_chars=2000,
                new_after_n_chars=6000
            )
        elif new_file_path.endswith(".docx"):
            return partition_docx(
                filename=new_file_path,
                strategy="hi_res",
                infer_table_structure=True,
                include_page_breaks=True,
                extract_image_block_types=["Image"],
                extract_image_block_to_payload=True,
                chunking_strategy="by_title",
                max_characters=10000,
                combine_text_under_n_chars=2000,
                new_after_n_chars=6000
            )
    except Exception as e:
        st.error(f"Error during chunking: {e}")
        return []

# Function to process chunks and prepare them for indexing
def chunk_processing(chunks, file_name):
    document_list = []
    for id, chunk in enumerate(chunks):
        chunk_dict = chunk.to_dict()
        document_list.append(
            Document(
                page_content=chunk_dict['text'],
                metadata={
                    'chunk_id': f"{file_name}_{id}",
                    'filename': file_name,
                    'page_number': str(chunk_dict['metadata'].get('page_number', 'NA'))
                }
            )
        )
    return document_list

# Function to index documents with their embeddings into Elasticsearch
def index_documents_to_es(documents):
    for doc in documents:
        embedding_vector = embeddings_client.encode(doc.page_content)  # Generate embedding for the document content
        es.index(index=INDEX_NAME, body={
            "embedding": embedding_vector,
            "page_content": doc.page_content,
            "metadata": doc.metadata
        })

# Function to clear all uploaded files from a folder
def clear_folder(folder_path):
    try:
        if os.path.exists(folder_path):
            shutil.rmtree(folder_path)  # Remove the folder and its contents
            os.makedirs(folder_path, exist_ok=True)  # Recreate an empty folder
    except Exception as e:
        st.error(f"Error clearing folder: {e}")

# Function to retrieve all indexed documents from Elasticsearch
def get_all_items_from_index(size=100):
    try:
        response = es.search(index=INDEX_NAME, body={"query": {"match_all": {}}, "size": size})
        return [hit['_source'] for hit in response['hits']['hits']]  # Extract document data
    except Exception as e:
        st.error(f"Error retrieving items from index: {e}")
        return []

# Function to perform a k-NN search in Elasticsearch based on a query
def query_index(query, rsize=5):
    query_embedding = embeddings_client.encode(query)  # Convert query to embedding

    # Perform k-NN search using Elasticsearch's kNN feature
    response = es.search(
        index=INDEX_NAME,
        body={
            "query": {
                "bool": {
                    "should": [
                        {
                            "multi_match": {
                                "query": query,
                                "fields": ["page_content"],  # Adjust field names based on index
                                "boost": 0.3
                            }
                        },
                        {
                            "knn": {
                                "field": "embedding",  # The field where embeddings are stored
                                "query_vector": query_embedding,  # Query embedding
                                "k": rsize,  # Number of nearest neighbors to retrieve
                                "num_candidates": 100,  # Number of candidates to consider in the search
                                "boost": 0.7
                            }
                        }
                    ]
                }
            },
            "size": rsize
        }
    )
    return response['hits']['hits']

# Function to delete the Elasticsearch index and remove all uploaded files
def delete_index(index_name=INDEX_NAME):
    try:
        if es.indices.exists(index=index_name):
            es.indices.delete(index=index_name)  # Delete the index from Elasticsearch
            st.success("Index deleted and folders cleared successfully.")
        else:
            st.warning("Index does not exist.")
        clear_folder(TEMP_FOLDER)
        clear_folder(UPLOADS_FOLDER)
        st.success("Uploaded files removed, and uploads, temp folders cleared.")
    except Exception as e:
        st.error(f"Error deleting index: {e}")

# Streamlit app to interact with the document indexing and search system
def app():
    st.title("Document Indexing and Search")

    # File upload section
    uploaded_files = st.file_uploader("Upload PDF/DOCX files", accept_multiple_files=True)
    new_files_uploaded = False  # Track if new files are uploaded
    
    if uploaded_files:
        for uploaded_file in uploaded_files:
            if is_document_indexed(uploaded_file.name):
                st.write(f"{uploaded_file.name} is already indexed. Skipping.")
                continue  
            # Save uploaded doc in temp folder
            file_path, new_file_path = asyncio.run(save_uploaded_file(uploaded_file))
            chunks = document_retrieval_chunking(new_file_path)
            if chunks:
                document_list = chunk_processing(chunks, uploaded_file.name)
                index_documents_to_es(document_list)  # Index the documents into Elasticsearch
                st.write(f"{uploaded_file.name} uploaded, translated, chunked and indexed")

    # Delete index button (allows the user to delete the Elasticsearch index and clear uploaded files)
    st.header("Manage Index")
    if st.button("Delete Index"):
        delete_index()

    # Display all indexed documents
    st.header("All Indexed Documents")
    documents = get_all_items_from_index()
    if documents:
        for doc in documents:
            st.subheader(f"{doc['metadata']['filename']} (Chunk ID: {doc['metadata']['chunk_id']})")
            st.write(f"Content: {doc['page_content']}")  # Show a snippet of the content
            st.write("---")
    else:
        st.write("No documents found in the index.")

    # Search functionality
    st.header("Search the Index")
    user_query = st.text_input("Enter a search query:")
    if user_query:
        st.subheader("Search Results")
        search_results = query_index(user_query)  # Perform the search using the user's query
        if search_results:
            for result in search_results:
                doc = result['_source']
                st.write(f"{doc['metadata']['filename']} (Chunk ID: {doc['metadata']['chunk_id']})")
                st.write(((result["_score"] * 2) - 1))  # Display similarity distance
                st.write("---")
        else:
            st.write("No results found.")

if __name__ == "__main__":
    app()
